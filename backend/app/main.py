from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict
import io
import uvicorn
from app.document_retrieval import DocumentRetrieval
from app.chatbot import Chatbot
from app.document_management import DocumentManager
import os
import json
import uuid
from datetime import datetime
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
from app.utils.file_handler import validate_file, save_file, get_file_metadata
from app.utils.logger import setup_logger
from app.config import get_settings

app = FastAPI(
    title="Biotech Regulatory Compliance Tool API",
    description="""
    An API for managing and querying regulatory documents for biotech companies.
    This tool helps companies navigate complex regulatory requirements and ensure compliance.
    """,
    version="1.0.0",
    docs_url="/api/docs",
    redoc_url="/api/redoc",
    openapi_url="/api/openapi.json"
)

logger = setup_logger(__name__)
settings = get_settings()

# Initialize components
doc_retrieval = DocumentRetrieval()
chatbot = Chatbot(doc_retrieval)
doc_manager = DocumentManager()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.BACKEND_CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# Data models
class QuestionnaireInput(BaseModel):
    intended_purpose: str
    life_threatening: bool
    user_type: str
    requires_sterilization: bool
    body_contact_duration: str

class ChatMessage(BaseModel):
    message: str
    questionnaire_data: Optional[dict] = None

class ChatResponse(BaseModel):
    response: str
    sources: Optional[List[dict]] = None

class RegulatoryGuideline(BaseModel):
    title: str
    content: str
    reference: str
    relevance_score: float

class BulkUploadMetadata(BaseModel):
    metadata_list: List[Dict]

class BulkDeleteRequest(BaseModel):
    doc_ids: List[str]

@app.get("/")
async def root():
    """
    API root endpoint.
    
    Returns:
    - A welcome message
    
    Raises:
    - None
    """
    return {"message": "Biotech Regulatory Compliance Tool API"}

@app.post("/chat", response_model=ChatResponse)
async def chat(message: ChatMessage):
    """
    Get a response from the AI chatbot.
    
    Parameters:
    - message: The user's question about regulatory requirements
    - questionnaire_data: Optional JSON string containing questionnaire responses
    
    Returns:
    - AI-generated response based on the query and available documents
    
    Raises:
    - 400: Invalid questionnaire data format
    - 500: Server error during chat processing
    """
    try:
        # Get chatbot response with questionnaire context
        response = await chatbot.get_response(
            message.message,
            questionnaire_data=message.questionnaire_data
        )
        
        # Get relevant documents used as sources
        sources = doc_retrieval.search(message.message)
        
        # Get user documents
        try:
            with open("user_documents.json", "r") as f:
                user_docs = json.load(f)
        except FileNotFoundError:
            user_docs = []
            
        # Process user documents
        user_sources = []
        for doc in user_docs:
            try:
                file_path = doc["file_path"]
                if os.path.exists(file_path):
                    # Extract text based on file type
                    file_extension = os.path.splitext(file_path)[1].lower()
                    doc_text = ""
                    
                    if file_extension == '.pdf':
                        with open(file_path, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            for page in pdf_reader.pages:
                                doc_text += page.extract_text() + "\n"
                    elif file_extension in ['.doc', '.docx']:
                        doc_obj = Document(file_path)
                        doc_text = "\n".join([paragraph.text for paragraph in doc_obj.paragraphs])
                    else:
                        with open(file_path, 'r') as file:
                            doc_text = file.read()
                            
                    # Add to sources if relevant
                    relevance_score = doc_retrieval.calculate_relevance(message.message, doc_text)
                    if relevance_score > 0.3:  # Adjust threshold as needed
                        user_sources.append({
                            "title": doc["title"],
                            "content": doc_text[:200] + "..." if len(doc_text) > 200 else doc_text,
                            "jurisdiction": "User Document"
                        })
            except Exception as e:
                print(f"Error processing user document {doc.get('title', 'Unknown')}: {str(e)}")
                continue
        
        # Combine sources
        all_sources = sources + user_sources
        
        # Update response to mention user documents if they were used
        if user_sources:
            response += "\n\nThis response also considers information from your uploaded documents."
        
        return ChatResponse(
            response=response,
            sources=[{
                "title": source.get("title", "Unknown"),
                "content": source.get("content", "")[:200] + "..." if len(source.get("content", "")) > 200 else source.get("content", ""),
                "jurisdiction": source.get("jurisdiction", "Unknown")
            } for source in all_sources]
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/questionnaire")
async def process_questionnaire(input_data: QuestionnaireInput):
    """
    Process a regulatory questionnaire.
    
    Parameters:
    - input_data: Questionnaire input data
    
    Returns:
    - Success message and processed data
    
    Raises:
    - 500: Server error during questionnaire processing
    """
    try:
        # TODO: Implement document retrieval based on questionnaire input
        return {
            "status": "success",
            "message": "Questionnaire processed successfully",
            "data": input_data.dict()
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/guidelines")
async def get_guidelines(query: str):
    """
    Get regulatory guidelines based on a query.
    
    Parameters:
    - query: Search query for guidelines
    
    Returns:
    - List of relevant guidelines
    
    Raises:
    - 500: Server error during guideline retrieval
    """
    try:
        # TODO: Implement guideline retrieval logic
        sample_guideline = RegulatoryGuideline(
            title="Sample Guideline",
            content="This is a placeholder guideline content",
            reference="REF-001",
            relevance_score=0.95
        )
        return [sample_guideline]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Document Management Endpoints
@app.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    title: str = Form(...),
    document_type: str = Form(...),
    jurisdiction: str = Form(...),
    version: str = Form(...),
    effective_date: str = Form(...),
    categories: str = Form(""),
    tags: str = Form("")
):
    """
    Upload a new document with metadata.
    
    Parameters:
    - file: The document file (PDF, DOCX, or TXT)
    - title: Document title
    - document_type: Document type (e.g., "Regulatory Guideline", "Clinical Trial Protocol")
    - jurisdiction: Regulatory jurisdiction (e.g., "EU", "US", "Global")
    - version: Document version
    - effective_date: Effective date of the document
    - categories: Comma-separated categories (e.g., "Medical Devices", "Clinical Trials")
    - tags: Comma-separated tags
    
    Returns:
    - Success message and document metadata
    
    Raises:
    - 400: Invalid file type or size
    - 500: Server error during upload
    """
    try:
        categories_list = [c.strip() for c in categories.split(",")] if categories else []
        tags_list = [t.strip() for t in tags.split(",")] if tags else []
        
        result = await doc_manager.upload_document(
            file=file,
            title=title,
            document_type=document_type,
            jurisdiction=jurisdiction,
            version=version,
            effective_date=effective_date,
            categories=categories_list,
            tags=tags_list
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/documents/bulk-upload")
async def bulk_upload_documents(
    metadata: BulkUploadMetadata,
    files: List[UploadFile] = File([])
):
    """
    Upload multiple documents at once.
    
    Parameters:
    - metadata: List of document metadata
    - files: List of document files
    
    Returns:
    - Success message
    
    Raises:
    - 400: Number of files does not match number of metadata entries
    - 500: Server error during upload
    """
    if len(files) != len(metadata.metadata_list):
        raise HTTPException(
            status_code=400,
            detail="Number of files must match number of metadata entries"
        )
    return await doc_manager.bulk_upload(files, metadata.metadata_list)

@app.post("/documents/bulk-delete")
async def bulk_delete_documents(request: BulkDeleteRequest):
    """
    Delete multiple documents at once.
    
    Parameters:
    - request: List of document IDs to delete
    
    Returns:
    - Success message
    
    Raises:
    - 500: Server error during deletion
    """
    return doc_manager.bulk_delete(request.doc_ids)

@app.get("/documents")
async def list_documents(
    document_type: Optional[str] = None,
    jurisdiction: Optional[str] = None,
    category: Optional[str] = None,
    tag: Optional[str] = None
):
    """
    List all documents with optional filtering.
    
    Parameters:
    - document_type: Filter by document type
    - jurisdiction: Filter by jurisdiction
    - category: Filter by category
    - tag: Filter by tag
    
    Returns:
    - List of documents
    
    Raises:
    - 500: Server error during document retrieval
    """
    return doc_manager.list_documents(
        document_type=document_type,
        jurisdiction=jurisdiction,
        category=category,
        tag=tag
    )

@app.get("/documents/{doc_id}")
async def get_document(doc_id: str):
    """
    Get metadata for a specific document.
    
    Parameters:
    - doc_id: ID of the document to retrieve
    
    Returns:
    - Document metadata
    
    Raises:
    - 404: Document not found
    - 500: Server error during document retrieval
    """
    result = doc_manager.get_document_metadata(doc_id)
    if not result:
        raise HTTPException(status_code=404, detail="Document not found")
    return result

@app.get("/documents/{doc_id}/preview")
async def get_document_preview(doc_id: str):
    """
    Get document preview content.
    
    Parameters:
    - doc_id: ID of the document to retrieve
    
    Returns:
    - Document preview content
    
    Raises:
    - 404: Document not found
    - 500: Server error during document retrieval
    """
    return doc_manager.get_document_preview(doc_id)

@app.get("/documents/{doc_id}/versions")
async def get_document_versions(doc_id: str):
    """
    Get version history for a document.
    
    Parameters:
    - doc_id: ID of the document to retrieve
    
    Returns:
    - List of document versions
    
    Raises:
    - 404: Document not found
    - 500: Server error during document retrieval
    """
    return doc_manager.get_document_versions(doc_id)

@app.get("/documents/categories")
async def get_categories():
    """
    Get all unique categories.
    
    Returns:
    - List of categories
    
    Raises:
    - 500: Server error during category retrieval
    """
    return doc_manager.get_categories()

@app.get("/documents/tags")
async def get_tags():
    """
    Get all unique tags.
    
    Returns:
    - List of tags
    
    Raises:
    - 500: Server error during tag retrieval
    """
    return doc_manager.get_tags()

@app.get("/documents/export")
async def export_documents(format: str = Query("json", regex="^(json|csv)$")):
    """
    Export document list in various formats.
    
    Parameters:
    - format: Export format ("json" or "csv")
    
    Returns:
    - Exported document list
    
    Raises:
    - 400: Invalid export format
    - 500: Server error during export
    """
    try:
        content = doc_manager.export_document_list(format)
        
        if format == "json":
            media_type = "application/json"
            filename = "documents.json"
        else:  # csv
            media_type = "text/csv"
            filename = "documents.csv"
        
        return StreamingResponse(
            io.StringIO(content),
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """
    Delete a document.
    
    Parameters:
    - doc_id: ID of the document to delete
    
    Returns:
    - Success message
    
    Raises:
    - 404: Document not found
    - 500: Server error during deletion
    """
    try:
        doc_manager.delete_document(doc_id)
        return {"message": "Document deleted successfully"}
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.patch("/documents/{doc_id}")
async def update_document(doc_id: str, updates: Dict):
    """
    Update document metadata.
    
    Parameters:
    - doc_id: ID of the document to update
    - updates: Dictionary of metadata updates
    
    Returns:
    - Updated document metadata
    
    Raises:
    - 404: Document not found
    - 500: Server error during update
    """
    try:
        result = doc_manager.update_document_metadata(doc_id, updates)
        return result
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# User Document Management
@app.post("/user-documents/upload")
async def upload_user_document(
    file: UploadFile,
    title: str = Form(None),
    description: str = Form("")
):
    """
    Upload a user document.
    
    Parameters:
    - file: The document file (PDF, DOCX, or TXT)
    - title: Document title
    - description: Document description
    
    Returns:
    - Success message and document ID
    
    Raises:
    - 400: Invalid file type or size
    - 500: Server error during upload
    """
    try:
        # Validate file
        content_type, content = validate_file(file.file, file.filename)
        
        # Save file
        file_path = save_file(content, file.filename, "user_documents")
        
        # Get file metadata
        metadata = get_file_metadata(file_path)
        
        # Process document content
        doc_retrieval = DocumentRetrieval()
        doc_id = doc_retrieval.index_document({
            "content": content.decode('utf-8'),
            "metadata": {
                "filename": file.filename,
                "content_type": content_type,
                **metadata
            }
        })
        
        logger.info(f"Document uploaded and processed successfully: {doc_id}")
        return {"message": "File uploaded successfully", "document_id": doc_id}
        
    except Exception as e:
        logger.error(f"Error processing upload: {str(e)}")
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/user-documents")
async def get_user_documents():
    """
    Get all user documents.
    
    Returns:
    - List of user documents
    
    Raises:
    - 500: Server error during document retrieval
    """
    try:
        with open("user_documents.json", "r") as f:
            user_docs = json.load(f)
        return user_docs
    except FileNotFoundError:
        return []
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/user-documents/{doc_id}/preview")
async def get_user_document_preview(doc_id: str):
    """
    Get user document preview content.
    
    Parameters:
    - doc_id: ID of the document to retrieve
    
    Returns:
    - Document preview content
    
    Raises:
    - 404: Document not found
    - 500: Server error during document retrieval
    """
    try:
        # Get document metadata
        with open("user_documents.json", "r") as f:
            user_docs = json.load(f)
        
        doc = next((d for d in user_docs if d["id"] == doc_id), None)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Read file content
        file_path = doc["file_path"]
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")

        # Extract text based on file type
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif file_extension in ['.doc', '.docx']:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            text = "Preview not available for this file type"

        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/user-documents/{doc_id}")
async def delete_user_document(doc_id: str):
    """
    Delete a user document.
    
    Parameters:
    - doc_id: ID of the document to delete
    
    Returns:
    - Success message
    
    Raises:
    - 404: Document not found
    - 500: Server error during deletion
    """
    try:
        # Read current documents
        with open("user_documents.json", "r") as f:
            user_docs = json.load(f)
        
        # Find document
        doc = next((d for d in user_docs if d["id"] == doc_id), None)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete file
        file_path = doc["file_path"]
        if os.path.exists(file_path):
            os.remove(file_path)

        # Update database
        user_docs = [d for d in user_docs if d["id"] != doc_id]
        with open("user_documents.json", "w") as f:
            json.dump(user_docs, f, indent=2)

        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    print("Starting server...")
    uvicorn.run(
        "app.main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="debug"
    )
